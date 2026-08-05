from __future__ import annotations

import csv
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
THESIS = ROOT / "thesis"
FIGURES = THESIS / "figures"
TABLES = THESIS / "tables"
BUILD = THESIS / "build"


def ensure_dirs() -> None:
    for path in (FIGURES, TABLES, BUILD):
        path.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill, outline=(30, 45, 65)) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font(20, True)) <= (x2 - x1 - 28):
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    total_h = len(lines) * 25
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=font(20, True))
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill=(20, 30, 40), font=font(20, True))
        y += 25


def arrow(draw: ImageDraw.ImageDraw, start, end, color=(35, 60, 90)) -> None:
    draw.line([start, end], fill=color, width=3)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        pts = [(ex, ey), (ex - 12, ey - 7), (ex - 12, ey + 7)]
    else:
        pts = [(ex, ey), (ex + 12, ey - 7), (ex + 12, ey + 7)]
    draw.polygon(pts, fill=color)


def make_architecture_figure() -> Path:
    out = FIGURES / "figure_5_1_system_architecture.png"
    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((45, 35), "System Architecture of the Adaptive Fusion Framework", fill=(20, 30, 40), font=font(34, True))

    boxes = [
        ((70, 150, 330, 260), "Fakeddit Metadata and Images", (222, 236, 255)),
        ((430, 90, 700, 200), "Image Branch ResNet", (229, 245, 224)),
        ((430, 250, 700, 360), "Text Branch DistilBERT", (255, 237, 213)),
        ((800, 170, 1070, 290), "Reliability-Aware State", (242, 232, 255)),
        ((1160, 170, 1430, 290), "RL Adaptive Fusion Policy", (255, 228, 230)),
        ((800, 430, 1070, 540), "Baselines and Ablations", (239, 246, 255)),
        ((1160, 430, 1430, 540), "Evaluation and Result Export", (240, 253, 244)),
        ((800, 620, 1070, 730), "Grad-CAM and Token Saliency", (254, 249, 195)),
        ((1160, 620, 1430, 730), "Explainable Final Prediction", (236, 253, 245)),
    ]
    for box in boxes:
        draw_box(d, *box)

    arrow(d, (330, 205), (430, 145))
    arrow(d, (330, 205), (430, 305))
    arrow(d, (700, 145), (800, 215))
    arrow(d, (700, 305), (800, 245))
    arrow(d, (1070, 230), (1160, 230))
    arrow(d, (935, 290), (935, 430))
    arrow(d, (1070, 485), (1160, 485))
    arrow(d, (935, 540), (935, 620))
    arrow(d, (1070, 675), (1160, 675))
    img.save(out)
    return out


def make_workflow_figure() -> Path:
    out = FIGURES / "figure_4_1_research_workflow.png"
    img = Image.new("RGB", (1500, 720), "white")
    d = ImageDraw.Draw(img)
    d.text((45, 35), "Research Workflow for RL-Based Adaptive Image-Text Fusion", fill=(20, 30, 40), font=font(34, True))
    labels = [
        "Prepare Fakeddit Data",
        "Download Available Images",
        "Train Image and Text Branches",
        "Build Reliability State",
        "Train Fusion Policy",
        "Evaluate, Explain, Export",
    ]
    x = 80
    prev = None
    for i, label in enumerate(labels):
        xy = (x, 250, x + 210, 380)
        draw_box(d, xy, f"{i + 1}. {label}", (235, 245, 255))
        if prev:
            arrow(d, (prev + 210, 315), (x, 315))
        prev = x
        x += 235
    d.text((80, 500), "Validation data was used for model selection and threshold tuning. Test data was used only for final reporting.", fill=(65, 75, 90), font=font(24))
    img.save(out)
    return out


def make_bar_chart(values: list[tuple[str, float]], title: str, out: Path, x_min: float = 0.75) -> Path:
    img = Image.new("RGB", (1200, 720), "white")
    d = ImageDraw.Draw(img)
    d.text((45, 30), title, fill=(20, 30, 40), font=font(32, True))
    left, top, width, bar_h, gap = 360, 120, 710, 38, 22
    max_v = max(v for _, v in values)
    span = max_v - x_min
    for i, (name, value) in enumerate(values):
        y = top + i * (bar_h + gap)
        d.text((45, y + 4), name, fill=(30, 45, 65), font=font(21))
        bar_w = int(((value - x_min) / span) * width) if span > 0 else width
        d.rounded_rectangle((left, y, left + max(4, bar_w), y + bar_h), radius=8, fill=(54, 116, 181))
        d.text((left + bar_w + 14, y + 3), f"{value:.4f}", fill=(20, 30, 40), font=font(21, True))
    d.line((left, top - 15, left, top + len(values) * (bar_h + gap) - gap + 12), fill=(180, 190, 200), width=2)
    img.save(out)
    return out


def make_robustness_figure() -> Path:
    out = FIGURES / "figure_7_4_robustness_image_quality.png"
    df = pd.read_csv(ROOT / "outputs/robustness/final_robustness_summary.csv")
    df = df[df["corruption"] == "image_quality_drop"].sort_values("level")
    img = Image.new("RGB", (1100, 680), "white")
    d = ImageDraw.Draw(img)
    d.text((45, 30), "Image Quality Degradation and Image Weight Adaptation", fill=(20, 30, 40), font=font(30, True))
    x0, y0, w, h = 120, 550, 820, 380
    d.line((x0, y0, x0 + w, y0), fill=(40, 50, 60), width=2)
    d.line((x0, y0, x0, y0 - h), fill=(40, 50, 60), width=2)
    levels = list(df["level"])
    vals = list(df["delta_target_weight"])
    min_v, max_v = -0.14, 0.0
    pts = []
    for level, val in zip(levels, vals):
        x = x0 + int((level - 0.25) / 0.5 * w)
        y = y0 - int((val - min_v) / (max_v - min_v) * h)
        pts.append((x, y))
        d.ellipse((x - 8, y - 8, x + 8, y + 8), fill=(200, 60, 70))
        d.text((x - 25, y - 42), f"{val:.4f}", fill=(20, 30, 40), font=font(18, True))
        d.text((x - 20, y0 + 18), f"{level:.2f}", fill=(20, 30, 40), font=font(18))
    d.line(pts, fill=(200, 60, 70), width=4)
    d.text((410, 610), "Corruption level", fill=(20, 30, 40), font=font(21, True))
    d.text((30, 250), "Image-weight change", fill=(20, 30, 40), font=font(21, True))
    img.save(out)
    return out


def export_tables() -> list[Path]:
    outputs = []
    method = pd.read_csv(ROOT / "outputs/tables/final_method_comparison.csv")
    method_test = method[method["split"] == "test"].copy()
    method_test = method_test[["source", "method", "macro_f1", "accuracy", "roc_auc"]]
    out = TABLES / "table_7_1_final_method_comparison.csv"
    method_test.to_csv(out, index=False)
    outputs.append(out)

    tuned = pd.read_csv(ROOT / "outputs/metrics/final_threshold_tuning_matched_summary.csv")
    out = TABLES / "table_7_2_threshold_tuning.csv"
    tuned.to_csv(out, index=False)
    outputs.append(out)

    seed = pd.read_csv(ROOT / "outputs/metrics/final_seed_significance_summary.csv")
    out = TABLES / "table_7_5_seed_stability.csv"
    seed.to_csv(out, index=False)
    outputs.append(out)
    return outputs


def create_figures() -> list[Path]:
    figs = [
        make_workflow_figure(),
        make_architecture_figure(),
    ]
    method = pd.read_csv(ROOT / "outputs/tables/final_method_comparison.csv")
    test = method[method["split"] == "test"]
    wanted = [
        "image_only",
        "text_only",
        "reliability_weighted_fusion",
        "equal_fusion",
        "rl_full_state",
        "supervised_mlp_fusion_matched",
    ]
    values = []
    for name in wanted:
        row = test[test["method"] == name]
        if not row.empty:
            values.append((name.replace("_", " "), float(row.iloc[0]["macro_f1"])))
    figs.append(make_bar_chart(values, "Final Test Macro F1 Comparison", FIGURES / "figure_7_1_method_macro_f1.png", 0.75))

    ab = pd.read_csv(ROOT / "outputs/metrics/final_ablation_summary.csv")
    vals = [(r["name"].replace("_", " "), float(r["test_macro_f1"])) for _, r in ab.iterrows()]
    figs.append(make_bar_chart(vals, "Ablation Results by Fusion State Representation", FIGURES / "figure_7_2_ablation_macro_f1.png", 0.84))
    figs.append(make_robustness_figure())
    return figs


def set_doc_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_title_page(doc: Document, metadata: dict) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(metadata["title"])
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"
    for text in [
        metadata["candidate_name"],
        metadata["registration_number"],
        "Thesis submitted in partial fulfillment of the requirements for the degree Research Project in Artificial Intelligence for the degree of BSc Hons in Artificial Intelligence",
        metadata["department"],
        metadata["faculty"],
        metadata["university"],
        metadata["country"],
        metadata["submission_month_year"],
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    doc.add_page_break()


def add_markdown_file(doc: Document, path: Path, page_break_before: bool = True) -> None:
    if page_break_before:
        doc.add_page_break()
    lines = path.read_text(encoding="utf-8").splitlines()
    in_table = []

    def flush_table():
        nonlocal in_table
        if len(in_table) >= 2:
            rows = []
            for line in in_table:
                if re.match(r"^\|\s*-", line):
                    continue
                cells = [c.strip() for c in line.strip("|").split("|")]
                rows.append(cells)
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                table.style = "Table Grid"
                for j, val in enumerate(rows[0]):
                    table.rows[0].cells[j].text = val
                for row in rows[1:]:
                    cells = table.add_row().cells
                    for j, val in enumerate(row[: len(cells)]):
                        cells[j].text = val
                doc.add_paragraph()
        in_table = []

    for line in lines:
        if line.strip().startswith("|"):
            in_table.append(line)
            continue
        flush_table()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(stripped, style=None)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            para = doc.add_paragraph(stripped)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    flush_table()


def add_figures_to_doc(doc: Document, figures: list[Path]) -> None:
    captions = {
        "figure_4_1_research_workflow.png": "Figure 4.1: Research workflow for the proposed adaptive fusion framework.",
        "figure_5_1_system_architecture.png": "Figure 5.1: System architecture of the adaptive fusion framework.",
        "figure_7_1_method_macro_f1.png": "Figure 7.1: Final test macro F1 comparison across selected methods.",
        "figure_7_2_ablation_macro_f1.png": "Figure 7.2: Ablation results for different RL fusion state representations.",
        "figure_7_4_robustness_image_quality.png": "Figure 7.4: Image quality degradation and image-weight adaptation.",
    }
    doc.add_page_break()
    doc.add_heading("Generated Figures", level=1)
    for fig in figures:
        doc.add_picture(str(fig), width=Inches(6.3))
        p = doc.add_paragraph(captions.get(fig.name, fig.name))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc: Document, fig: Path, caption: str) -> None:
    doc.add_picture(str(fig), width=Inches(6.3))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_csv_table(doc: Document, csv_path: Path, caption: str, max_rows: int | None = None) -> None:
    rows = list(csv.reader(csv_path.open(encoding="utf-8")))
    if max_rows is not None and len(rows) > max_rows + 1:
        rows = rows[: max_rows + 1]
    doc.add_paragraph(caption)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for j, value in enumerate(rows[0]):
        table.rows[0].cells[j].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(row[: len(cells)]):
            cells[j].text = value
    doc.add_paragraph()


def build_docx(figures: list[Path]) -> Path:
    metadata = {}
    for line in (THESIS / "metadata.yaml").read_text(encoding="utf-8").splitlines():
        if ":" in line and not line.startswith(" ") and not line.startswith("-"):
            key, val = line.split(":", 1)
            metadata[key] = val.strip().strip('"')

    doc = Document()
    set_doc_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    add_page_number_footer(doc)

    add_title_page(doc, metadata)
    for fm in ["declaration.md", "dedication.md", "acknowledgement.md", "abstract_draft.md"]:
        add_markdown_file(doc, THESIS / "front_matter" / fm, page_break_before=True)

    add_markdown_file(doc, THESIS / "front_matter" / "table_of_contents_draft.md", page_break_before=True)
    for fm in ["list_of_figures.md", "list_of_tables.md", "list_of_abbreviations.md", "list_of_appendices.md"]:
        add_markdown_file(doc, THESIS / "front_matter" / fm, page_break_before=True)

    figure_map = {p.name: p for p in figures}
    for chapter in sorted((THESIS / "chapters").glob("*.md")):
        add_markdown_file(doc, chapter, page_break_before=True)
        if chapter.name == "04_approach.md":
            add_figure(doc, figure_map["figure_4_1_research_workflow.png"], "Figure 4.1: Research workflow for the proposed adaptive fusion framework.")
        elif chapter.name == "05_analysis_design.md":
            add_figure(doc, figure_map["figure_5_1_system_architecture.png"], "Figure 5.1: System architecture of the adaptive fusion framework.")
        elif chapter.name == "07_evaluation.md":
            add_csv_table(doc, TABLES / "table_7_1_final_method_comparison.csv", "Table 7.1: Final test method comparison.")
            add_figure(doc, figure_map["figure_7_1_method_macro_f1.png"], "Figure 7.1: Final test macro F1 comparison across selected methods.")
            add_csv_table(doc, TABLES / "table_7_2_threshold_tuning.csv", "Table 7.2: Validation-selected threshold tuning results.")
            add_figure(doc, figure_map["figure_7_2_ablation_macro_f1.png"], "Figure 7.2: Ablation results for different RL fusion state representations.")
            add_figure(doc, figure_map["figure_7_4_robustness_image_quality.png"], "Figure 7.4: Image quality degradation and image-weight adaptation.")

    add_markdown_file(doc, THESIS / "references" / "initial_ieee_references.md", page_break_before=True)
    add_markdown_file(doc, THESIS / "appendices" / "artifact_manifest.md", page_break_before=True)

    out = BUILD / "Nishanth_Thesis_Professor_Reviewed_Draft.docx"
    doc.save(out)
    legacy = BUILD / "Nishanth_Thesis_First_Draft.docx"
    doc.save(legacy)
    return out


def main() -> None:
    ensure_dirs()
    tables = export_tables()
    figures = create_figures()
    docx = build_docx(figures)
    summary = {
        "docx": str(docx),
        "figures": [str(p) for p in figures],
        "tables": [str(p) for p in tables],
    }
    (BUILD / "build_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
